from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from typing import List, Optional
import os
import tempfile
import json
import re
import zipfile
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from io import BytesIO

# Supabase
try:
    from supabase import create_client, Client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False

# Claude API
try:
    import anthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    CLAUDE_AVAILABLE = False

app = FastAPI(title="AI Doc Pro API")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Supabase client
def get_supabase() -> Optional[Client]:
    url = os.getenv("SUPABASE_URL")
    key = os.getenv("SUPABASE_SERVICE_KEY")
    if url and key and SUPABASE_AVAILABLE:
        return create_client(url, key)
    return None

# Claude client
def get_claude_client():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key or not CLAUDE_AVAILABLE:
        return None
    return anthropic.Anthropic(api_key=api_key)

# Pydantic models
class ExcelRequest(BaseModel):
    prompt: str

class CheckLimitRequest(BaseModel):
    user_id: Optional[str] = None

# ============ HELPERS ============

def get_client_ip(request: Request) -> str:
    """Client IP manzilini olish"""
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else "unknown"

async def check_and_record_usage(request: Request, action_type: str, user_id: Optional[str] = None) -> dict:
    """Limitni tekshirish va foydalanishni qayd qilish"""
    supabase = get_supabase()
    if not supabase:
        # Supabase yo'q bo'lsa, ruxsat berish
        return {"allowed": True, "remaining": -1, "is_premium": False}
    
    ip_address = get_client_ip(request)
    
    try:
        # Function chaqirish
        result = supabase.rpc('record_usage', {
            'p_ip_address': ip_address,
            'p_action_type': action_type,
            'p_user_id': user_id
        }).execute()
        
        if result.data:
            return result.data
        return {"allowed": True, "remaining": -1}
    except Exception as e:
        print(f"Usage tracking error: {e}")
        return {"allowed": True, "remaining": -1}

async def check_limit_only(request: Request, user_id: Optional[str] = None) -> dict:
    """Faqat limitni tekshirish (qayd qilmasdan)"""
    supabase = get_supabase()
    if not supabase:
        return {"allowed": True, "remaining": -1, "is_premium": False}
    
    ip_address = get_client_ip(request)
    
    try:
        result = supabase.rpc('check_daily_limit', {
            'check_ip': ip_address,
            'check_user_id': user_id
        }).execute()
        
        if result.data:
            return result.data
        return {"allowed": True, "remaining": 5}
    except Exception as e:
        print(f"Limit check error: {e}")
        return {"allowed": True, "remaining": 5}

# ============ ENDPOINTS ============

@app.get("/")
async def root():
    return {
        "message": "AI Doc Pro API ishlamoqda!",
        "version": "2.0",
        "features": ["excel", "autofill", "templates", "auth"]
    }

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.post("/api/check-limit")
async def check_limit(request: Request, user_id: Optional[str] = None):
    """Foydalanuvchi limitini tekshirish"""
    result = await check_limit_only(request, user_id)
    return result

# ============ AI EXCEL GENERATOR ============

async def generate_excel_with_ai(prompt: str) -> dict:
    """Claude AI bilan Excel yaratish"""
    client = get_claude_client()
    
    if not client:
        return generate_excel_fallback(prompt)
    
    system_prompt = """Sen Excel jadval yaratuvchi AI assistantsan. Foydalanuvchi so'rovi asosida JSON formatda jadval strukturasini yarat.

MUHIM: Faqat JSON qaytar, boshqa hech narsa yo'q!

JSON formati:
{
    "title": "Fayl_nomi",
    "sheets": [{
        "name": "Varaq nomi",
        "headers": ["Ustun1", "Ustun2", "Ustun3"],
        "data": [
            ["qiymat1", "qiymat2", "qiymat3"],
            ["qiymat4", "qiymat5", "=A2+B2"]
        ]
    }]
}

Qoidalar:
1. Formulalar = bilan boshlanadi
2. Sonlar raqam sifatida
3. O'zbek tilida"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": f"Excel jadval yarat: {prompt}"}],
            system=system_prompt
        )
        
        response_text = message.content[0].text.strip()
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
        raise ValueError("JSON topilmadi")
    except Exception as e:
        print(f"AI error: {e}")
        return generate_excel_fallback(prompt)

def generate_excel_fallback(prompt: str) -> dict:
    """Fallback generator"""
    return {
        "title": "Jadval",
        "sheets": [{
            "name": "Ma'lumotlar",
            "headers": ["№", "Nomi", "Miqdori", "Narxi", "Jami"],
            "data": [
                [1, "Element 1", 10, 50000, "=C2*D2"],
                [2, "Element 2", 20, 30000, "=C3*D3"],
                ["", "", "", "JAMI:", "=SUM(E2:E3)"],
            ]
        }]
    }

def create_styled_excel(structure: dict) -> str:
    """Excel fayl yaratish"""
    wb = Workbook()
    
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    for sheet_idx, sheet_data in enumerate(structure.get("sheets", [])):
        ws = wb.active if sheet_idx == 0 else wb.create_sheet()
        ws.title = sheet_data.get("name", "Sheet1")[:31]
        
        headers = sheet_data.get("headers", [])
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
        
        for row_idx, row_data in enumerate(sheet_data.get("data", []), 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center")
        
        for col_idx in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = 18
    
    temp_dir = tempfile.mkdtemp()
    filepath = os.path.join(temp_dir, f"{structure.get('title', 'document')}.xlsx")
    wb.save(filepath)
    return filepath

@app.post("/api/excel/preview")
async def excel_preview(request: Request, data: ExcelRequest, user_id: Optional[str] = Header(None, alias="X-User-ID")):
    """Excel preview"""
    try:
        structure = await generate_excel_with_ai(data.prompt)
        return {"success": True, "title": structure.get("title"), "sheets": structure.get("sheets")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/excel/generate")
async def excel_generate(request: Request, data: ExcelRequest, user_id: Optional[str] = Header(None, alias="X-User-ID")):
    """Excel yaratish"""
    try:
        # Limitni tekshirish va qayd qilish
        usage_result = await check_and_record_usage(request, "excel", user_id)
        if not usage_result.get("success", True) and not usage_result.get("allowed", True):
            raise HTTPException(status_code=429, detail="Kunlik limit tugadi. Premium obunaga o'ting!")
        
        structure = await generate_excel_with_ai(data.prompt)
        filepath = create_styled_excel(structure)
        
        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=f"{structure.get('title', 'document')}.xlsx"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============ AUTO-FILL (KO'P FAYL) ============

def extract_text_from_docx(content: bytes) -> str:
    """Word fayldan matn olish"""
    doc = Document(BytesIO(content))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

async def get_replacements_from_ai(text: str, instruction: str) -> list:
    """AI dan almashtirish ro'yxatini olish"""
    client = get_claude_client()
    
    if not client:
        raise HTTPException(status_code=400, detail="AI xizmati mavjud emas")
    
    system_prompt = """Sen hujjat tahrirlovchi AI assistantsan.

VAZIFANG:
1. Hujjat matnini DIQQAT BILAN o'qi
2. Ko'rsatma asosida BARCHA almashtirilishi kerak bo'lgan joylarni top
3. JSON formatda almashtirish ro'yxatini qaytar

MUHIM:
- Bir xil ma'lumot BIR NECHA MARTA takrorlanishi mumkin - BARCHASINI top!
- "old" maydoni AYNAN hujjatdagi matn bo'lishi kerak

JSON formati:
{
    "replacements": [
        {"old": "eski matn", "new": "yangi matn"}
    ]
}

FAQAT JSON QAYTAR!"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            messages=[{
                "role": "user",
                "content": f"Hujjat:\n{text[:8000]}\n\nKo'rsatma:\n{instruction}\n\nJSON:"
            }],
            system=system_prompt
        )
        
        response_text = message.content[0].text.strip()
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            data = json.loads(json_match.group())
            return data.get("replacements", [])
        return []
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI xatolik: {str(e)}")

def apply_replacements_to_docx(content: bytes, replacements: list) -> bytes:
    """Word faylga almashtirishlar qo'llash"""
    doc = Document(BytesIO(content))
    
    def replace_in_paragraph(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return True
            # Murakkab holat
            if old_text in paragraph.text:
                full_text = paragraph.text.replace(old_text, new_text)
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    return True
        return False
    
    for repl in replacements:
        old_text = repl.get("old", "")
        new_text = repl.get("new", "")
        if not old_text:
            continue
        
        for para in doc.paragraphs:
            replace_in_paragraph(para, old_text, new_text or "")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, old_text, new_text or "")
        
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    replace_in_paragraph(para, old_text, new_text or "")
            if section.footer:
                for para in section.footer.paragraphs:
                    replace_in_paragraph(para, old_text, new_text or "")
    
    # BytesIO ga saqlash
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()

@app.post("/api/autofill/process")
async def process_autofill(
    request: Request,
    files: List[UploadFile] = File(...),
    instruction: str = Form(...),
    user_id: Optional[str] = Header(None, alias="X-User-ID")
):
    """Ko'p faylni AI bilan tahrirlash"""
    try:
        if not instruction.strip():
            raise HTTPException(status_code=400, detail="Ko'rsatma kiriting")
        
        if len(files) == 0:
            raise HTTPException(status_code=400, detail="Fayl yuklang")
        
        if len(files) > 10:
            raise HTTPException(status_code=400, detail="Maksimum 10 ta fayl yuklash mumkin")
        
        # Limitni tekshirish
        usage_result = await check_and_record_usage(request, "autofill", user_id)
        if not usage_result.get("success", True) and not usage_result.get("allowed", True):
            raise HTTPException(status_code=429, detail="Kunlik limit tugadi. Premium obunaga o'ting!")
        
        processed_files = []
        
        for file in files:
            # Faqat Word qabul qilish
            if not file.filename.lower().endswith('.docx'):
                raise HTTPException(
                    status_code=400, 
                    detail=f"Faqat Word (.docx) fayllar qabul qilinadi. '{file.filename}' qabul qilinmadi."
                )
            
            content = await file.read()
            
            # Matnni olish
            text = extract_text_from_docx(content)
            if not text.strip():
                continue
            
            # AI dan almashtirishlar olish
            replacements = await get_replacements_from_ai(text, instruction)
            
            if replacements:
                # Almashtirishlarni qo'llash
                modified_content = apply_replacements_to_docx(content, replacements)
                processed_files.append({
                    "filename": f"tahrirlangan_{file.filename}",
                    "content": modified_content
                })
            else:
                processed_files.append({
                    "filename": f"tahrirlangan_{file.filename}",
                    "content": content
                })
        
        if len(processed_files) == 0:
            raise HTTPException(status_code=400, detail="Hech qanday fayl qayta ishlana olmadi")
        
        # Bitta fayl bo'lsa
        if len(processed_files) == 1:
            temp_dir = tempfile.mkdtemp()
            output_path = os.path.join(temp_dir, processed_files[0]["filename"])
            with open(output_path, 'wb') as f:
                f.write(processed_files[0]["content"])
            
            return FileResponse(
                output_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=processed_files[0]["filename"]
            )
        
        # Ko'p fayl bo'lsa - ZIP
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, "tahrirlangan_hujjatlar.zip")
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for pf in processed_files:
                zf.writestr(pf["filename"], pf["content"])
        
        return FileResponse(
            zip_path,
            media_type="application/zip",
            filename="tahrirlangan_hujjatlar.zip"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Server xatoligi: {str(e)}")

# ============ TEMPLATES ============

@app.get("/api/templates")
async def get_templates(request: Request, user_id: Optional[str] = Header(None, alias="X-User-ID")):
    """Shablonlar ro'yxatini olish"""
    supabase = get_supabase()
    if not supabase:
        return {"templates": [], "message": "Database mavjud emas"}
    
    try:
        # Public shablonlar
        result = supabase.table('templates').select('*').eq('is_public', True).execute()
        templates = result.data or []
        
        # User shablonlari
        if user_id:
            user_result = supabase.table('templates').select('*').eq('user_id', user_id).execute()
            templates.extend(user_result.data or [])
        
        return {"templates": templates}
    except Exception as e:
        return {"templates": [], "error": str(e)}

@app.post("/api/templates")
async def create_template(
    request: Request,
    file: UploadFile = File(...),
    name: str = Form(...),
    description: str = Form(""),
    category: str = Form("other"),
    user_id: str = Header(..., alias="X-User-ID")
):
    """Yangi shablon yaratish"""
    supabase = get_supabase()
    if not supabase:
        raise HTTPException(status_code=500, detail="Database mavjud emas")
    
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Faqat Word (.docx) fayllar qabul qilinadi")
    
    try:
        content = await file.read()
        
        # Faylni storage ga yuklash
        file_path = f"templates/{user_id}/{file.filename}"
        supabase.storage.from_('templates').upload(file_path, content)
        
        # URL olish
        file_url = supabase.storage.from_('templates').get_public_url(file_path)
        
        # Database ga yozish
        result = supabase.table('templates').insert({
            'user_id': user_id,
            'name': name,
            'description': description,
            'category': category,
            'file_url': file_url,
            'file_name': file.filename,
            'is_public': False
        }).execute()
        
        return {"success": True, "template": result.data[0] if result.data else None}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/templates/{template_id}")
async def delete_template(
    template_id: str,
    user_id: str = Header(..., alias="X-User-ID")
):
    """Shablonni o'chirish"""
    supabase = get_supabase()
    if not supabase:
        raise HTTPException(status_code=500, detail="Database mavjud emas")
    
    try:
        supabase.table('templates').delete().eq('id', template_id).eq('user_id', user_id).execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ============ SUBSCRIPTION ============

@app.get("/api/subscription")
async def get_subscription(user_id: str = Header(..., alias="X-User-ID")):
    """Foydalanuvchi obunasini olish"""
    supabase = get_supabase()
    if not supabase:
        return {"plan": "free", "status": "active"}
    
    try:
        result = supabase.table('subscriptions').select('*').eq('user_id', user_id).order('created_at', desc=True).limit(1).execute()
        if result.data:
            return result.data[0]
        return {"plan": "free", "status": "active"}
    except Exception as e:
        return {"plan": "free", "status": "active", "error": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
